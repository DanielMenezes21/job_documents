from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from app.contrato_hono.PJ.texto_clausula import TEXTOS_CLAUSULAS
from docx import Document
import os
from app.contrato_hono.PJ.extracao_PJ_contrato import formatar_data , substituir_palavras_documento
from app.homepage.home_screen import HomeScreen

def voltar(self, instance):
    """
    Função para voltar para a tela anterior ou a tela principal.
    """
    self.manager.current = "contrato_screen1"
    
def atualizar_dados(screen_instance, dados):
    """
    Atualiza os dados recebidos na tela Processo2Screen.
    """
    screen_instance.dados = dados
    screen_instance.caminho_modelo = dados.get("caminho_modelo", None)  # Armazena o caminho do arquivo modelo
    print(f"Dados recebidos: {screen_instance.dados}")

def on_spinner_select(self, spinner, text):
    """
    Atualiza o TextInput com o texto da cláusula selecionada e salva o texto atual no dicionário de placeholders.
    """
    # Salva o texto atual da cláusula previamente selecionada
    if spinner.text in TEXTOS_CLAUSULAS:
        selected_clause_index = list(TEXTOS_CLAUSULAS.keys()).index(spinner.text) + 1
        self.placeholders[f"#CLAUSULA{selected_clause_index}"] = self.text_input.text

    # Atualiza o TextInput com o texto da nova cláusula selecionada
    self.text_input.text = TEXTOS_CLAUSULAS.get(text, "")
    
def on_text_change(self, instance, value):
    """
    Atualiza o dicionário de placeholders sempre que o texto no TextInput for alterado.
    """
    if self.spinner.text in TEXTOS_CLAUSULAS:
        selected_clause_index = list(TEXTOS_CLAUSULAS.keys()).index(self.spinner.text) + 1
        self.placeholders[f"#CLAUSULA{selected_clause_index}"] = value
        
def salvar_texto(self, screen_instance):
    """
    Salva o texto editado e atualiza o documento.
    """
    try:
        
        selected_clause_index = list(TEXTOS_CLAUSULAS.keys()).index(self.spinner.text) + 1
        self.placeholders[f"#CLAUSULA{selected_clause_index}"] = self.text_input.text
        
        if not screen_instance.dados:
            mostrar_popup(screen_instance, "Erro", "Nenhum dado foi recebido da tela inicial!")
            return

        if not screen_instance.caminho_modelo:
            mostrar_popup(screen_instance, "Erro", "O arquivo modelo não foi selecionado na tela inicial.")
            return

        # Verifica se o arquivo existe
        if not os.path.exists(screen_instance.caminho_modelo):
            mostrar_popup(screen_instance, "Erro", f"O arquivo {screen_instance.caminho_modelo} não foi encontrado!")
            return
        for idx, key in enumerate(TEXTOS_CLAUSULAS.keys(), start=1):
            if f"#CLAUSULA{idx}" not in self.placeholders:
                self.placeholders[f"#CLAUSULA{idx}"] = TEXTOS_CLAUSULAS[key]

        document = Document(screen_instance.caminho_modelo)
        data_atual = formatar_data()

        # Adiciona os placeholders restantes
        placeholders = self.placeholders.copy()
        placeholders.update({
            "#NOME_CLIENTE": screen_instance.dados.get("nome_cliente", ''),
            "#NME_EMPRESA": screen_instance.dados.get("nome_empresa", ''),
            "#CNPJ":screen_instance.dados.get("cnpj", ''),
            "#NUMERO": screen_instance.dados.get("numero", ''),
            "#END_EMPRESA": screen_instance.dados.get("end_empresa", ''),
            "#CEP_EMPRESA": screen_instance.dados.get("cep_empresa", ''),
            "#NACIONALIDADE": screen_instance.dados.get("nacionalidade", ''),
            "#ESTADO_CIVIL": screen_instance.dados.get("estado_civil", ''),
            "#END_CLIENTE": screen_instance.dados.get("endereco_cliente", ''),
            "#CEP_CLIENTE": screen_instance.dados.get("cep_cliente", ''),
            "#CPF": screen_instance.dados.get("cpf", ''),
            "#RG": screen_instance.dados.get("rg", ''),
            "#SEC_RG": screen_instance.dados.get("sec_rg", ''),
            "#CIDADE_EMP": screen_instance.dados.get("cidade_emp", ''),
            "#ESTADO_EMP": screen_instance.dados.get("estado_emp", ''),
            "#CIDADE_cliente": screen_instance.dados.get("cidade_cliente", ''),
            "#SIGLA_ESTADO_CLIENTE": screen_instance.dados.get("sigla_estado_cliente", ''),
            "#INSCRITA(O)": screen_instance.dados.get("inscrita_o", ''),
            "#DATA_AGORA": data_atual,
        })
        
        print(f"mostrar {placeholders}")

        # Função para substituir os placeholders mantendo a formatação
        def substituir_com_formatacao(paragrafo, placeholders):
            for run in paragrafo.runs:
                for placeholder, valor in placeholders.items():
                    if valor is None:
                        valor = ''
                        print(f"o placeholder {placeholder} está vazio")
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, valor)
                        #print(f"substituindo {placeholder} por {valor}")

        # Substituir os placeholders no documento
        for paragraph in document.paragraphs:
            substituir_com_formatacao(paragraph, placeholders)

        # Substituir nas tabelas
        for tabela in document.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        substituir_com_formatacao(paragrafo, placeholders)

        # Salvar o documento final com o nome especificado
        nome_arquivo = screen_instance.dados.get("nome_arquivo", "documento_final")  # Usar nome_arquivo, se disponível
        caminho_salvamento = f"{nome_arquivo}.docx"
        document.save(caminho_salvamento)
        mostrar_popup(screen_instance, "Sucesso", f"Documento salvo em {caminho_salvamento}")
        os.startfile(caminho_salvamento)
            
    except Exception as e:
        print(f"Erro ao obter dados: {e}")
        return None
    

def mostrar_popup(screen_instance, titulo, mensagem):
    """
    Exibe um popup com uma mensagem.
    """
    conteudo = BoxLayout(orientation='vertical', padding=10, spacing=10)
    conteudo.add_widget(Label(text=mensagem, halign='center'))
    btn_fechar = Button(text='Fechar', size_hint=(1, 0.3))
    popup = Popup(title=titulo, content=conteudo, size_hint=(0.7, 0.4))
    btn_fechar.bind(on_press=popup.dismiss)
    conteudo.add_widget(btn_fechar)
    popup.open()