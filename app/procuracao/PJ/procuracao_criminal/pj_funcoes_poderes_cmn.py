from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from app.procuracao.PJ.procuracao_criminal.pj_texto_poderes_cmn import TEXTOS_PODERES_A, ADVOGADO_OAB
from app.procuracao.PJ.procuracao_criminal.pj_extracao_procuracao_cmn import *
from app.homepage.home_screen import *
from docx import Document
import sqlite3
import os
    
def voltar(self, instance):
    """
 Função para voltar para a tela anterior ou a tela principal.
    """
    self.manager.current = "procuracao_criminal_pj_screen"

def atualizar_dados(screen_instance, dados):
    """
    Atualiza os dados recebidos na tela PoderesScreen.
    """
    screen_instance.dados = dados
    screen_instance.caminho_modelo = dados.get("caminho_modelo", None)  
    screen_instance.caminho_declaracao = dados.get("caminho_declaracao", None)  
    screen_instance.dados["advogado_id"] = screen_instance.manager.advogado_id
    print(f"Dados recebidos: {screen_instance.dados}")

def on_text_selected(self, modelo_spinner, text):
    """
    Carrega o texto selecionado no TextInput.
    """
    if text in TEXTOS_PODERES_A:
        self.text_input.text = TEXTOS_PODERES_A[text]

def obter_dados_advogado(advogado_id):
    """
    Obtém os dados do advogado a partir do banco de dados.
    """
    conn = sqlite3.connect('advogados.db')
    cursor = conn.cursor()
    cursor.execute("SELECT username, identidade_oab FROM advogados WHERE id = ?", (advogado_id,))
    resultado = cursor.fetchone()
    conn.close()
    if resultado:
        print(f"nome {resultado[0]} oab {resultado[1]}")
        return {"nome": resultado[0], "oab": resultado[1]}
    else:
        return {"nome": "Advogado não encontrado", "oab": "OAB não encontrado"}

def salvar_texto(screen_instance, _):
    """
    Salva o texto editado e atualiza o documento.
    """
    if not screen_instance.dados:
        mostrar_popup(screen_instance, "Erro", "Nenhum dado foi recebido da tela inicial!")
        return

    texto_editado = screen_instance.text_input.text

    if not screen_instance.caminho_modelo:
        mostrar_popup(screen_instance, "Erro", "O arquivo modelo não foi selecionado na tela inicial.")
        return

    if not os.path.exists(screen_instance.caminho_modelo):
        mostrar_popup(screen_instance, "Erro", f"O arquivo {screen_instance.caminho_modelo} não foi encontrado!")
        return
    
    if not os.path.exists(screen_instance.caminho_declaracao):
        mostrar_popup(screen_instance, "Erro", f"o arquivo{screen_instance.caminho_declaracao} não foi encontrado!", placeholders={}, caminho_declaracao="None")
        return

    document = Document(screen_instance.caminho_modelo)

    advogado_id = screen_instance.dados.get("advogado_id")
    dados_advogado = obter_dados_advogado(advogado_id)

    placeholders = {
        "#PODERES": texto_editado,
        "#NOME_EMPRESA": screen_instance.dados.get("nome_empresa"),
        "#CNPJ":screen_instance.dados.get("cnpj"),
        "#END_EMPRESA": screen_instance.dados.get("end_empresa"),
        "#CP_EMPRESA": screen_instance.dados.get("cep_empresa"),
        "#CIDADE_EMPRESA": screen_instance.dados.get("cidade_empresa"),
        "#ESTADO_EMPRESA": screen_instance.dados.get("estado_empresa"),
        "#SIGLA_ESTADO_CLIENTE": screen_instance.dados.get("estado_empresa"),
        "#NOME_CLIENTE": screen_instance.dados.get("nome_cliente"),
        "#CPF": screen_instance.dados.get("cpf"),
        "#CIDADE_CLIENTE": screen_instance.dados.get("cidade_cliente"),
        "#SIGLA_ESTADO_CLIENTE": screen_instance.dados.get("sigla_estado_cliente"),
        "#INSCRITA(O)": screen_instance.dados.get("inscrita_o"),
        "#NACIONALIDADE": screen_instance.dados.get("nacionalidade"),
        "#NOME_ARQUIVO": screen_instance.dados.get("nome_arquivo"),
        "#DATA_AGORA": screen_instance.dados.get("data_agora"),
        "#RG_CLIENTE": screen_instance.dados.get("rg"),
        "#SEC_RG": screen_instance.dados.get("sec_rg"),
        "#EST_RG": screen_instance.dados.get("est_rg"),
        "#ENDERECO": screen_instance.dados.get("endereco"),
        "#CEP": screen_instance.dados.get("cep"),
        "#ESTADO_CIVIL": screen_instance.dados.get("estado_civil"),
        "#ADVOGADO_OAB": dados_advogado["oab"],
        "#ADVOGADO_NOME": dados_advogado["nome"]
    }
    
    print(f"mostrar {placeholders}")

    def substituir_com_formatacao(paragrafo, placeholders):
        for run in paragrafo.runs:
            for placeholder, valor in placeholders.items():
                if valor is None:
                    valor = ''
                    print(f"o placeholder {placeholder} está vazio")
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)

    for paragraph in document.paragraphs:
        substituir_com_formatacao(paragraph, placeholders)
        print(f"substituindo {placeholders} no documento")

    for tabela in document.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_com_formatacao(paragrafo, placeholders)


    nome_arquivo = screen_instance.dados.get("nome_arquivo", "documento_final") 
    caminho_salvamento = f"{nome_arquivo}.docx"
    document.save(caminho_salvamento)
    mostrar_popup(
        screen_instance,
        "Sucesso",
        f"Documento salvo como {nome_arquivo}. Deseja gerar uma declaração?",
        placeholders,
        screen_instance.dados.get("caminho_declaracao")  
    )
    os.startfile(caminho_salvamento)

def mostrar_popup(screen_instance, titulo, mensagem, placeholders, caminho_declaracao):
    conteudo = BoxLayout(orientation='vertical', padding=10, spacing=10)
    conteudo.add_widget(Label(text=mensagem, halign='center'))

    text_input = TextInput(hint_text="Digite o nome do arquivo", multiline=False)
    conteudo.add_widget(text_input)

    btn_fechar = Button(text='Fechar', size_hint=(1, 0.3))
    btn_fechar.bind(on_press=lambda instance: popup.dismiss())

    btn_gerar_declaracao = Button(text='Gerar Declaração', size_hint=(1, 0.3))

    def gerar_declaracao(instance):
        if not os.path.exists(caminho_declaracao):
            print("Modelo de declaração não encontrado.")
            popup.dismiss()
            return
        try:
            document = Document(caminho_declaracao)
            for paragraph in document.paragraphs:
                substituir_com_formatacao(paragraph, placeholders)
            for tabela in document.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            substituir_com_formatacao(paragrafo, placeholders)

            nome_arquivo = text_input.text.strip() or "declaracao_gerada"
            cd_arquivo = f"{nome_arquivo}.docx"
            document.save(cd_arquivo)
            print(f"Declaração gerada e salva como {cd_arquivo}.")
            os.startfile(cd_arquivo)
        except Exception as e:
            print(f"Erro ao gerar declaração: {e}")
        popup.dismiss()

    btn_gerar_declaracao.bind(on_press=gerar_declaracao)

    conteudo.add_widget(btn_fechar)
    conteudo.add_widget(btn_gerar_declaracao)

    popup = Popup(title=titulo, content=conteudo, size_hint=(0.7, 0.5))
    popup.open()