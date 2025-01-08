#funcoes_poderes.py
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import Screen
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from kivy.uix.widget import Widget
from kivy.uix.floatlayout import FloatLayout
from kivy.graphics import Ellipse, Color
from telas.procuracao.procuracao_queixa_crime.texto_poderes import TEXTOS_PODERES, ADVOGADO_OAB
from telas.procuracao.procuracao_queixa_crime.extracao_procuracao_qc import processar_documento, formatar_data

from telas.homepage.home_screen import *
from docx import Document
import os
    
def voltar(self, instance):
    """
 Função para voltar para a tela anterior ou a tela principal.
    """
    self.manager.current = "procuracao_screen"

def atualizar_dados(screen_instance, dados):
    """
    Atualiza os dados recebidos na tela PoderesScreen.
    """
    screen_instance.dados = dados
    screen_instance.caminho_modelo = dados.get("caminho_modelo", None)  # Armazena o caminho do arquivo modelo
    print(f"Dados recebidos: {screen_instance.dados}")

def on_text_selected(self, modelo_spinner, text):
    """
    Carrega o texto selecionado no TextInput.
    """
    if text in TEXTOS_PODERES:
        self.text_input.text = TEXTOS_PODERES[text]  # Preenche o TextInput com o texto selecionado

def salvar_texto(screen_instance, _):
    """
    Salva o texto editado e atualiza o documento.
    """
    if not screen_instance.dados:
        mostrar_popup(screen_instance, "Erro", "Nenhum dado foi recebido da tela inicial!")
        return

    # Obtém o texto editado
    texto_editado = screen_instance.text_input.text

    if not screen_instance.caminho_modelo:
        mostrar_popup(screen_instance, "Erro", "O arquivo modelo não foi selecionado na tela inicial.")
        return

    # Verifica se o arquivo existe
    if not os.path.exists(screen_instance.caminho_modelo):
        mostrar_popup(screen_instance, "Erro", f"O arquivo {screen_instance.caminho_modelo} não foi encontrado!")
        return

    # Abrir o arquivo modelo
    document = Document(screen_instance.caminho_modelo)

    # Adiciona o advogado OAB baseado no nome selecionado no spinner
    advogado_nome = screen_instance.adv_spinner.text
    advogado_oab = ADVOGADO_OAB.get(advogado_nome, "OAB não encontrado")  # Default: "OAB não encontrado"

    # Dicionário de placeholders com o texto editado
    placeholders = {
        "#PODERES": texto_editado,
        "#NOME_OUTORGANTE": screen_instance.dados.get("nome_outorgante"),
        "#OUTORGANTE_CPF": screen_instance.dados.get("cpf"),
        "#CIDADE_OUTORGANTE": screen_instance.dados.get("cidade_outorgante"),
        "#SIGLA_ESTADO_OUTORGANTE": screen_instance.dados.get("sigla_estado_outorgante"),
        "#INSCRITA(O)": screen_instance.dados.get("inscrita_o"),
        "#NACIONALIDADE": screen_instance.dados.get("nacionalidade"),
        "#NOME_ARQUIVO": screen_instance.dados.get("nome_arquivo"),
        "#DATA_AGORA": screen_instance.dados.get("data_agora"),
        "#RG_OUTORGANTE": screen_instance.dados.get("rg"),
        "#ENDERECO": screen_instance.dados.get("endereco"),
        "#CEP": screen_instance.dados.get("cep"),
        "#ESTADO_CIVIL": screen_instance.dados.get("estado_civil"),
        "#ADVOGADO_OAB": advogado_oab,
    }
    
    print(f"mostrar {placeholders}")

    # Função para substituir os placeholders mantendo a formatação
    def substituir_com_formatacao(paragrafo, placeholders):
        for run in paragrafo.runs:
            for placeholder, valor in placeholders.items():
                if valor is None:
                    valor = ''
                    print(f"o placeholder {placeholder} está vazio")
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)
                    #print(f"substituindo {placeholder} por {valor}")

    # Substituir os placeholders no documento
    for paragraph in document.paragraphs:
        substituir_com_formatacao(paragraph, placeholders)
        print(f"substituindo {placeholders} no documento")

    # Substituir nas tabelas
    for tabela in document.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_com_formatacao(paragrafo, placeholders)

    # Salvar o documento final com o nome especificado
    nome_arquivo = screen_instance.dados.get("nome_arquivo", "documento_final")  # Usar nome_arquivo, se disponível
    caminho_salvamento = f"{nome_arquivo}.docx"
    document.save(caminho_salvamento)
    mostrar_popup(screen_instance, "Sucesso", f"Documento salvo em {caminho_salvamento}")

def mostrar_popup(screen_instance, titulo, mensagem):
    """
    Exibe um popup com uma mensagem.
    """
    conteudo = BoxLayout(orientation='vertical', padding=10, spacing=10)
    conteudo.add_widget(Label(text=mensagem, halign='center'))
    btn_fechar = Button(text='Fechar', size_hint=(1, 0.3))
    popup = Popup(title=titulo, content=conteudo, size_hint=(0.7, 0.4))
    btn_fechar.bind(on_press=popup.dismiss)
    conteudo.add_widget(btn_fechar)
    popup.open()