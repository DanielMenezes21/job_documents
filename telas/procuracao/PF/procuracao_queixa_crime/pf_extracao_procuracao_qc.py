#exracao_procuracao.py
from docx import Document
from datetime import datetime

def formatar_data(data):
    """
    Formata uma data do tipo datetime no formato 'dia de mês de ano'.
    """
    if not isinstance(data, datetime):
        raise ValueError("O parâmetro 'data' deve ser um objeto datetime.")
    
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }
    
    dia = data.day
    mes = meses.get(data.month, "mês inválido")  # Garantir retorno válido mesmo em caso de erro
    ano = data.year
    
    data_formatada = f"{dia} de {mes} de {ano}"
    return data_formatada


def substituir_com_formatacao(paragrafo, placeholders):
    """
    Substitui os placeholders no parágrafo mantendo a formatação original.
    """
    for run in paragrafo.runs:
        for placeholder, valor in placeholders.items():
            if valor is None:
                valor = ''
            if placeholder in run.text:
                # Exibe qual placeholder foi encontrado e o texto original
                print(f"Substituindo: {placeholder} por {valor} em {run.text}")
                # Substitui o texto mantendo a formatação do run
                run.text = run.text.replace(placeholder, valor)

def processar_documento(caminho_modelo, nome_outorgante, cpf, cidade_outorgante, 
                        sigla_estado_outorgante, inscrita_o, nacionalidade, poderes, 
                        nome_arquivo, advogado_oab, endereço, cep, estado_civil, rg):
    # Abrir o arquivo modelo
    document = Document(caminho_modelo)

    # Obter a data atual formatada
    data_atual = datetime.now()
    print(f"mostrando: {data_atual}")
    data_agora = formatar_data(data_atual)
    print(f"Data formatada: {data_agora}")
    
    # Dicionário com os marcadores e seus valores correspondentes
    placeholders = {
        "#NOME_OUTORGANTE": nome_outorgante,
        "#OUTORGANTE_CPF": cpf,
        "#CIDADE_OUTORGANTE": cidade_outorgante,
        "#SIGLA_ESTADO_OUTORGANTE": sigla_estado_outorgante,
        "#DATA_AGORA": data_agora,
        "#INSCRITA(O)": inscrita_o,
        "#NACIONALIDADE": nacionalidade,
        "#PODERES": poderes,
        "#ADVOGADO_OAB": advogado_oab,
        "#RG_OUTORGANTE": rg,
        "#ENDERECO": endereço,
        "#CEP": cep,
        "#ESTADO_CIVIL": estado_civil
    }

    for placeholder in placeholders:
        print(f"Verificando placeholder: {placeholder}")
        if placeholder not in [run.text for paragraph in document.paragraphs for run in paragraph.runs]:
            print(f"Atenção: {placeholder} não encontrado no documento!")

    # Substituir os placeholders no documento, mantendo a formatação
    for paragraph in document.paragraphs:
        substituir_com_formatacao(paragraph, placeholders)

    # Substituindo nas tabelas
    for tabela in document.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_com_formatacao(paragrafo, placeholders)

    # Salvar o documento final com o nome especificado
    caminho_salvamento = f"{nome_arquivo}.docx"
    document.save(caminho_salvamento)

