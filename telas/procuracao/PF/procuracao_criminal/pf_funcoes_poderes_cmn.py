#funcoes_poderes.py
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import Screen
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from kivy.uix.widget import Widget
from kivy.uix.floatlayout import FloatLayout
from kivy.graphics import Ellipse, Color
from telas.procuracao.PF.procuracao_criminal.pf_texto_poderes_cmn import TEXTOS_PODERES_A, ADVOGADO_OAB
from telas.procuracao.PF.procuracao_criminal.pf_extracao_procuracao_cmn import processar_documento, formatar_data, substituir_com_formatacao
from telas.homepage.home_screen import *
from docx import Document
import os
    
def voltar(self, instance):
    """
 Função para voltar para a tela anterior ou a tela principal.
    """
    self.manager.current = "procuracao_criminal_screen_PF"

def atualizar_dados(screen_instance, dados):
    """
    Atualiza os dados recebidos na tela PoderesScreen.
    """
    screen_instance.dados = dados
    screen_instance.caminho_modelo = dados.get("caminho_modelo", None)  # Armazena o caminho do arquivo modelo
    screen_instance.caminho_declaracao = dados.get("caminho_declaracao", None)  # Armazena o caminho do arquivo modelo da declaração
    print(f"Dados recebidos: {screen_instance.dados}")

def on_text_selected(self, modelo_spinner, text):
    """
    Carrega o texto selecionado no TextInput.
    """
    if text in TEXTOS_PODERES_A:
        self.text_input.text = TEXTOS_PODERES_A[text]  # Preenche o TextInput com o texto selecionado

def salvar_texto(screen_instance, _):
    """
    Salva o texto editado e atualiza o documento.
    """
    if not screen_instance.dados:
        mostrar_popup(screen_instance, "Erro", "Nenhum dado foi recebido da tela inicial!")
        return

    # Obtém o texto editado
    texto_editado = screen_instance.text_input.text

    if not screen_instance.caminho_modelo:
        mostrar_popup(screen_instance, "Erro", "O arquivo modelo não foi selecionado na tela inicial.")
        return

    # Verifica se o arquivo existe
    if not os.path.exists(screen_instance.caminho_modelo):
        mostrar_popup(screen_instance, "Erro", f"O arquivo {screen_instance.caminho_modelo} não foi encontrado!")
        return
    
    if not os.path.exists(screen_instance.caminho_declaracao):
        mostrar_popup(screen_instance, "Erro", f"o arquivo {screen_instance.dados.get('caminho_declaracao')} não foi encontrado!", {}, None)

        return

    # Abrir o arquivo modelo
    document = Document(screen_instance.caminho_modelo)

    # Adiciona o advogado OAB baseado no nome selecionado no spinner
    advogado_nome = screen_instance.adv_spinner.text
    advogado_oab = ADVOGADO_OAB.get(advogado_nome, "OAB não encontrado")  # Default: "OAB não encontrado"

    # Dicionário de placeholders com o texto editado
    placeholders = {
        "#PODERES": texto_editado,
        "#NOME_OUTORGANTE": screen_instance.dados.get("nome_outorgante"),
        "#PROFISSAO": screen_instance.dados.get("profissao"),
        "#SEC_RG": screen_instance.dados.get("sec_rg"),
        "#EST_RG": screen_instance.dados.get("est_rg"),
        "#OUTORGANTE_CPF": screen_instance.dados.get("cpf"),
        "#CIDADE_OUTORGANTE": screen_instance.dados.get("cidade_outorgante"),
        "#SIGLA_ESTADO_OUTORGANTE": screen_instance.dados.get("sigla_estado_outorgante"),
        "#INSCRITA(O)": screen_instance.dados.get("inscrita_o"),
        "#NACIONALIDADE": screen_instance.dados.get("nacionalidade"),
        "#NOME_ARQUIVO": screen_instance.dados.get("nome_arquivo"),
        "#DATA_AGORA": screen_instance.dados.get("data_agora"),
        "#RG_OUTORGANTE": screen_instance.dados.get("rg"),
        "#ENDERECO": screen_instance.dados.get("endereco"),
        "#CEP": screen_instance.dados.get("cep"),
        "#ESTADO_CIVIL": screen_instance.dados.get("estado_civil"),
        "#ADVOGADO_OAB": advogado_oab,
    }
    
    print(f"mostrar {placeholders}")

    # Função para substituir os placeholders mantendo a formatação
    def substituir_com_formatacao(paragrafo, placeholders):
        for run in paragrafo.runs:
            for placeholder, valor in placeholders.items():
                if valor is None:
                    valor = ''
                    print(f"o placeholder {placeholder} está vazio")
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)
                    #print(f"substituindo {placeholder} por {valor}")

    # Substituir os placeholders no documento
    for paragraph in document.paragraphs:
        substituir_com_formatacao(paragraph, placeholders)
        print(f"substituindo {placeholders} no documento")

    # Substituir nas tabelas
    for tabela in document.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_com_formatacao(paragrafo, placeholders)

    # Salvar o documento final com o nome especificado
    nome_arquivo = screen_instance.dados.get("nome_arquivo", "documento_final")  # Usar nome_arquivo, se disponível
    caminho_salvamento = f"{nome_arquivo}.docx"
    document.save(caminho_salvamento)
    mostrar_popup(
        screen_instance,
        "Sucesso",
        f"Documento salvo como {nome_arquivo}. Deseja gerar uma declaração?",
        placeholders,
        screen_instance.dados.get("caminho_declaracao")  # Aqui você passa o caminho do modelo da declaração
    )
    os.startfile(caminho_salvamento)

def mostrar_popup(screen_instance, titulo, mensagem, placeholders=None, caminho_declaracao=None):
    if placeholders is None:
        placeholders = {}  # Defina um dicionário vazio se não for passado um valor
    if caminho_declaracao is None:
        caminho_declaracao = ""  # Defina uma string vazia se não for passado um valor

    conteudo = BoxLayout(orientation='vertical', padding=10, spacing=10)
    conteudo.add_widget(Label(text=mensagem, halign='center'))

    # Campo de entrada para o nome do arquivo
    text_input = TextInput(hint_text="Digite o nome do arquivo", multiline=False)
    conteudo.add_widget(text_input)

    btn_fechar = Button(text='Fechar', size_hint=(1, 0.3))
    btn_fechar.bind(on_press=lambda instance: popup.dismiss())

    btn_gerar_declaracao = Button(text='Gerar Declaração', size_hint=(1, 0.3))

    def gerar_declaracao(instance):
        if not os.path.exists(caminho_declaracao):
            print("Modelo de declaração não encontrado.")
            popup.dismiss()
            return

        try:
            document = Document(caminho_declaracao)
            for paragraph in document.paragraphs:
                substituir_com_formatacao(paragraph, placeholders)
            for tabela in document.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            substituir_com_formatacao(paragrafo, placeholders)

            # Pega o nome do arquivo do TextInput
            nome_arquivo = text_input.text.strip() or "declaracao_gerada"
            cd_arquivo = f"{nome_arquivo}.docx"
            document.save(cd_arquivo)
            print(f"Declaração gerada e salva como {cd_arquivo}.")
            os.startfile(cd_arquivo)
        except Exception as e:
            print(f"Erro ao gerar declaração: {e}")
        popup.dismiss()

    btn_gerar_declaracao.bind(on_press=gerar_declaracao)

    conteudo.add_widget(btn_fechar)
    conteudo.add_widget(btn_gerar_declaracao)

    popup = Popup(title=titulo, content=conteudo, size_hint=(0.7, 0.5))
    popup.open()

