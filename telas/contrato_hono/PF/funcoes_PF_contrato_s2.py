from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from telas.contrato_hono.PF.texto_clausula import TEXTOS_CLAUSULAS
from docx import Document
import os
from telas.contrato_hono.PF.extracao_PF_contrato import formatar_data , substituir_palavras_documento
from telas.homepage.home_screen import HomeScreen

def voltar(self, instance):
    """
    Função para voltar para a tela anterior ou a tela principal.
    """
    self.manager.current = "contrato_PF_screen1"
    
def atualizar_dados(screen_instance, dados):
    """
    Atualiza os dados recebidos na tela Processo2Screen.
    """
    screen_instance.dados = dados
    screen_instance.caminho_modelo = dados.get("caminho_modelo", None)  # Armazena o caminho do arquivo modelo
    print(f"Dados recebidos: {screen_instance.dados}")

def on_spinner_select(self, spinner, text):
    clause_key = f"Cláusula {list(TEXTOS_CLAUSULAS.keys()).index(text) + 1}"
    self.placeholders[f"#CLAUSULA{list(TEXTOS_CLAUSULAS.keys()).index(text) + 1}"] = self.text_input.text
    self.text_input.text = TEXTOS_CLAUSULAS.get(text, "")

# Função para salvar texto e atualizar o documento
def salvar_texto(self, screen_instance):
    try:
        # Atualiza os placeholders com os textos das cláusulas
        for idx, key in enumerate(self.placeholders.keys(), start=1):
            clausula_key = f"Cláusula {idx}"
            self.placeholders[key] = TEXTOS_CLAUSULAS.get(clausula_key, "")

        if not screen_instance.dados:
            mostrar_popup(screen_instance, "Erro", "Nenhum dado foi recebido da tela inicial!")
            return

        if not screen_instance.caminho_modelo:
            mostrar_popup(screen_instance, "Erro", "O arquivo modelo não foi selecionado na tela inicial.")
            return

        # Verifica se o arquivo existe
        if not os.path.exists(screen_instance.caminho_modelo):
            mostrar_popup(screen_instance, "Erro", f"O arquivo {screen_instance.caminho_modelo} não foi encontrado!")
            return

        document = Document(screen_instance.caminho_modelo)
        data_atual = formatar_data()  # Obter a data atual

        # Dicionário de placeholders
        placeholders = {
            "#CLAUSULA1": self.text_input.text,  # Aqui vai o texto da cláusula editada
            "#NOME_CONTRATANTE": screen_instance.dados.get("nome_contratante", ''),
            "#NUMERO": screen_instance.dados.get("numero", ''),
            "#NACIONALIDADE": screen_instance.dados.get("nacionalidade", ''),
            "#ESTADO_CIVIL": screen_instance.dados.get("estado_civil", ''),
            "#PROFISSAO": screen_instance.dados.get('profissao', ''),
            "#END_CONTRATANTE": screen_instance.dados.get("endereco_contratante", ''),
            "#CEP_CONTRATANTE": screen_instance.dados.get("cep_contratante", ''),
            "#CPF": screen_instance.dados.get("cpf", ''),
            "#RG": screen_instance.dados.get("rg", ''),
            "#SEC_RG": screen_instance.dados.get("sec_rg", ''),
            "#CIDADE_CONTRATANTE": screen_instance.dados.get("cidade_contratante", ''),
            "#SIGLA_ESTADO_CONTRATANTE": screen_instance.dados.get("sigla_estado_contratante", ''),
            "#INSCRITA(O)": screen_instance.dados.get("inscrita_o", ''),
            "#NUM_TEL": screen_instance.dados.get("telefone", ''),
            "#EMAIL": screen_instance.dados.get("email", ''),
            "#DATA_AGORA": data_atual,
        }

        print(f"placeholders {placeholders}")

        # Função para substituir os placeholders mantendo a formatação
        def substituir_com_formatacao(paragrafo, placeholders):
            for run in paragrafo.runs:
                for placeholder, valor in placeholders.items():
                    if valor is None:
                        valor = ''
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, valor)

        # Substituir os placeholders no documento
        for paragraph in document.paragraphs:
            substituir_com_formatacao(paragraph, placeholders)

        # Substituir nas tabelas
        for tabela in document.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        substituir_com_formatacao(paragrafo, placeholders)

        # Salvar o documento final
        nome_arquivo = screen_instance.dados.get("nome_arquivo", "documento_final")
        caminho_salvamento = f"{nome_arquivo}.docx"
        document.save(caminho_salvamento)
        mostrar_popup(screen_instance, "Sucesso", f"Documento salvo em {caminho_salvamento}")
        os.startfile(caminho_salvamento)

    except Exception as e:
        print(f"Erro ao obter dados: {e}")
        return None

# Função para exibir um popup
def mostrar_popup(screen_instance, titulo, mensagem):
    conteudo = BoxLayout(orientation='vertical', padding=10, spacing=10)
    conteudo.add_widget(Label(text=mensagem, halign='center'))
    btn_fechar = Button(text='Fechar', size_hint=(1, 0.3))
    popup = Popup(title=titulo, content=conteudo, size_hint=(0.7, 0.4))
    btn_fechar.bind(on_press=popup.dismiss)
    conteudo.add_widget(btn_fechar)
    popup.open()