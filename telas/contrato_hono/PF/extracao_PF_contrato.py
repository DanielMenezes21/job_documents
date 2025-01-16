from docx import Document
from datetime import datetime
from docx.shared import Pt
import time

def formatar_data(argumento='argumento'):
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    
    data = datetime.now()
    dia = data.day
    mes = meses[data.month - 1]
    ano = data.year
    
    return f"{dia} de {mes} de {ano}."

def substituir_palavras_documento(doc_path, substituicoes, output_path, delay=0.1):
    documento = Document(doc_path)
    
    for paragrafo in documento.paragraphs:
        for run in paragrafo.runs:
            for palavra_antiga, palavra_nova in substituicoes.items():
                if palavra_antiga in run.text:
                    run.text = run.text.replace(palavra_antiga, palavra_nova)
                    print(f"Substituindo '{palavra_antiga}' por '{palavra_nova}'")
                    time.sleep(delay)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for palavra_antiga, palavra_nova in substituicoes.items():
                            if palavra_antiga in run.text:
                                run.text = run.text.replace(palavra_antiga, palavra_nova)
                                print(f"Substituindo '{palavra_antiga}' por '{palavra_nova}'")
                                time.sleep(delay)


    documento.save(output_path)
    print(f"Substituição concluída. Documento salvo em {output_path}")
